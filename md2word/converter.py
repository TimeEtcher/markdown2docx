"""Markdown AST -> docx 转换核心。"""

from __future__ import annotations

import os
import re

import mistune
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

from .styles import apply_page_setup, apply_paragraph_format, style_run


# 引用标注：例如 [1]、[1-2]、[1,3]、[1-3,5]
CITATION_RE = re.compile(r"\[\d+(?:-\d+)?(?:,\s*\d+(?:-\d+)?)*\]")


class Converter:
    def __init__(self, config: dict):
        self.config = config
        self.body_cfg = config["body"]
        self.headings_cfg = config["headings"]
        self.code_cfg = config["code"]
        self.quote_cfg = config["quote"]
        self.md = mistune.create_markdown(renderer="ast", plugins=["math", "table"])
        self.base_dir = os.getcwd()  # 图片相对路径的基准目录，convert() 中更新
        self.ref_numbers: set[int] = set()
        self.in_references = False
        self.ref_section_level = 0
        self._bookmark_id = 0

    # ------------------------------------------------------------------ API
    def convert_file(self, md_path: str, output_path: str) -> str:
        with open(md_path, "r", encoding="utf-8") as f:
            text = f.read()
        self.base_dir = os.path.dirname(os.path.abspath(md_path))
        return self.convert_text(text, output_path)

    def convert_text(self, md_text: str, output_path: str) -> str:
        tokens = list(self.md(md_text))
        tokens = self._merge_citation_tokens(tokens)
        self._scan_references(tokens)
        self.in_references = False
        self.ref_section_level = 0
        self._bookmark_id = 0

        doc = Document()
        apply_page_setup(doc, self.config["page"])
        for token in tokens:
            self.render_block(doc, token, depth=0, in_quote=False)
        doc.save(output_path)
        return output_path

    # --------------------------------------------------------------- blocks
    def render_block(self, doc, token: dict, depth: int, in_quote: bool) -> None:
        ttype = token.get("type")
        handler = getattr(self, f"block_{ttype}", None)
        if handler is None:
            # 未知块级元素：尽量把其中的文本段落渲染出来
            for child in token.get("children", []) or []:
                self.render_block(doc, child, depth, in_quote)
            return
        handler(doc, token, depth, in_quote)

    def block_heading(self, doc, token, depth, in_quote):
        level = max(1, min(6, int(token.get("attrs", {}).get("level", 1))))
        cfg = self.headings_cfg.get(f"h{level}", {})
        # 用 Heading N 样式保留大纲级别（Word 导航窗格可用）；
        # run 级显式字体设置会覆盖样式自带的字体/颜色
        p = doc.add_paragraph(style=f"Heading {level}")

        heading_text = self._inline_text(token).strip()
        if heading_text in ("参考文献", "References"):
            self.in_references = True
            self.ref_section_level = level
        elif self.in_references and level <= self.ref_section_level:
            self.in_references = False

        # 段前/段后间距支持“空行”单位（以正文字号为基准换算成磅）
        base_font_size = float(self.body_cfg.get("font_size", 12))
        space_before = cfg.get("space_before")
        space_after = cfg.get("space_after")
        if cfg.get("space_before_lines") is not None:
            space_before = float(cfg["space_before_lines"]) * base_font_size
        if cfg.get("space_after_lines") is not None:
            space_after = float(cfg["space_after_lines"]) * base_font_size

        apply_paragraph_format(
            p,
            align=cfg.get("align"),
            line_spacing=cfg.get("line_spacing"),
            space_before=space_before,
            space_after=space_after,
        )
        self._apply_heading_number_sep(token, cfg)
        self.render_inlines(p, token.get("children", []), cfg)

    def block_paragraph(self, doc, token, depth, in_quote):
        p = doc.add_paragraph()
        cfg = self.body_cfg
        # 正文段前/段后间距支持“空行”单位（以正文字号为基准换算成磅）
        font_size = float(cfg.get("font_size", 12))
        space_before = cfg.get("space_before")
        space_after = cfg.get("space_after")
        if cfg.get("space_before_lines") is not None:
            space_before = float(cfg["space_before_lines"]) * font_size
        if cfg.get("space_after_lines") is not None:
            space_after = float(cfg["space_after_lines"]) * font_size

        apply_paragraph_format(
            p,
            line_spacing=cfg.get("line_spacing"),
            space_before=space_before,
            space_after=space_after,
            first_line_indent_chars=2 if (cfg.get("first_line_indent") and not in_quote) else 0,
            font_size=cfg.get("font_size"),
            left_indent_cm=self.quote_cfg.get("indent_cm") if in_quote else None,
        )
        self.render_inlines(p, token.get("children", []), cfg)
        if in_quote and self.quote_cfg.get("italic"):
            for run in p.runs:
                run.font.italic = True

    def block_block_text(self, doc, token, depth, in_quote):
        # 紧凑列表项内的段落
        self.block_paragraph(doc, token, depth, in_quote)

    def block_block_code(self, doc, token, depth, in_quote):
        cfg = self.code_cfg
        lines = str(token.get("raw", "")).rstrip("\n").split("\n")
        for i, line in enumerate(lines):
            p = doc.add_paragraph()
            apply_paragraph_format(p, line_spacing=cfg.get("line_spacing"),
                                   space_before=6 if i == 0 else 0,
                                   space_after=6 if i == len(lines) - 1 else 0)
            run = p.add_run(line if line else " ")
            style_run(run, cfg)

    def block_block_quote(self, doc, token, depth, in_quote):
        for child in token.get("children", []) or []:
            self.render_block(doc, child, depth, in_quote=True)

    def block_thematic_break(self, doc, token, depth, in_quote):
        p = doc.add_paragraph()
        apply_paragraph_format(p, align="center", space_before=6, space_after=6)
        run = p.add_run("─" * 30)
        style_run(run, self.body_cfg)

    def block_blank_line(self, doc, token, depth, in_quote):
        pass

    def block_list(self, doc, token, depth, in_quote):
        attrs = token.get("attrs", {})
        ordered = bool(attrs.get("ordered", token.get("ordered", False)))
        # 有序列表：每个列表创建独立编号实例并强制从 start 重新计数，
        # 避免 Word 中全文共用 List Number 序列导致编号连续累加
        num_id = None
        if ordered:
            start = int(attrs.get("start", 1) or 1)
            num_id = self._new_ordered_num(doc, start)
        for item in token.get("children", []) or []:
            if item.get("type") != "list_item":
                self.render_block(doc, item, depth, in_quote)
                continue
            self._render_list_item(doc, item, ordered, depth, in_quote, num_id)

    def _new_ordered_num(self, doc, start):
        """在 numbering.xml 中新建一个指向 ListNumber 编号定义的 w:num（带 startOverride）。"""
        from docx.oxml.ns import qn
        from lxml import etree

        # ListNumber 样式 -> numId -> abstractNumId
        styles_el = doc.styles.element
        style = styles_el.find(f"{qn('w:style')}[@{qn('w:styleId')}='ListNumber']")
        style_num_id = style.find(f"{qn('w:pPr')}/{qn('w:numPr')}/{qn('w:numId')}").get(qn("w:val"))
        numbering = doc.part.numbering_part.element
        abs_id = None
        max_num_id = 0
        for num in numbering.findall(qn("w:num")):
            nid = int(num.get(qn("w:numId")))
            max_num_id = max(max_num_id, nid)
            if str(nid) == style_num_id:
                abs_id = num.find(qn("w:abstractNumId")).get(qn("w:val"))
        new_id = max_num_id + 1
        num = etree.SubElement(numbering, qn("w:num"))
        num.set(qn("w:numId"), str(new_id))
        abs_el = etree.SubElement(num, qn("w:abstractNumId"))
        abs_el.set(qn("w:val"), abs_id)
        lvl_ov = etree.SubElement(num, qn("w:lvlOverride"))
        lvl_ov.set(qn("w:ilvl"), "0")
        start_ov = etree.SubElement(lvl_ov, qn("w:startOverride"))
        start_ov.set(qn("w:val"), str(start))
        return new_id

    @staticmethod
    def _set_num_pr(paragraph, num_id):
        """给段落显式指定编号实例（覆盖 List Number 样式自带的共享 numId）。"""
        from docx.oxml.ns import qn

        pPr = paragraph._p.get_or_add_pPr()
        numPr = pPr.get_or_add_numPr()
        ilvl = numPr.get_or_add_ilvl()
        ilvl.val = 0
        numId = numPr.get_or_add_numId()
        numId.val = num_id

    def _render_list_item(self, doc, item, ordered, depth, in_quote, num_id=None):
        if ordered:
            style_name = "List Number"
        else:
            style_name = "List Bullet"
            if depth > 0:
                style_name = f"{style_name} {min(depth + 1, 3)}"
        children = item.get("children", []) or []
        first_para_done = False
        for child in children:
            ctype = child.get("type")
            if ctype == "list":
                self.render_block(doc, child, depth + 1, in_quote)
            elif ctype in ("paragraph", "block_text"):
                p = doc.add_paragraph(style=style_name)
                apply_paragraph_format(
                    p, line_spacing=self.body_cfg.get("line_spacing"),
                    left_indent_cm=0.75 * depth if ordered and depth > 0 else None)
                if num_id is not None:
                    self._set_num_pr(p, num_id)
                self.render_inlines(p, child.get("children", []), self.body_cfg)
                first_para_done = True
            else:
                self.render_block(doc, child, depth, in_quote)
        if not first_para_done and not children:
            p = doc.add_paragraph(style=style_name)
            if num_id is not None:
                self._set_num_pr(p, num_id)
            style_run(p.add_run(""), self.body_cfg)

    def block_table(self, doc, token, depth, in_quote):
        head, body = None, []
        for child in token.get("children", []) or []:
            if child.get("type") == "table_head":
                head = child.get("children", []) or []
            elif child.get("type") == "table_body":
                body = child.get("children", []) or []
        ncols = len(head) if head else (len(body[0].get("children", [])) if body else 0)
        if not ncols:
            return
        tbl_cfg = self.config.get("table", {})
        nrows = (1 if head else 0) + len(body)
        table = doc.add_table(rows=nrows, cols=ncols)
        if str(tbl_cfg.get("style", "grid")).lower() == "three_line":
            self._apply_three_line_borders(table, has_head=bool(head))
        else:
            table.style = "Table Grid"
        header_align = tbl_cfg.get("header_align", "center")
        header_bold = bool(tbl_cfg.get("header_bold", True))
        row_idx = 0
        if head:
            for col_idx, cell in enumerate(head):
                self._fill_table_cell(table.rows[0].cells[col_idx], cell,
                                      bold=header_bold, align=header_align)
            row_idx = 1
        for row in body:
            cells = row.get("children", []) or []
            for col_idx, cell in enumerate(cells[:ncols]):
                self._fill_table_cell(table.rows[row_idx].cells[col_idx], cell,
                                      bold=False)
            row_idx += 1

    @staticmethod
    def _apply_three_line_borders(table, has_head):
        """三线表：上下粗线（1.5pt）+ 表头下细线（0.75pt），无竖线和内部横线。"""
        from docx.oxml.ns import qn
        from lxml import etree

        def border(parent, name, val, sz):
            el = etree.SubElement(parent, qn(f"w:{name}"))
            el.set(qn("w:val"), val)
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "auto")

        tblPr = table._tbl.tblPr
        borders = tblPr.makeelement(qn("w:tblBorders"), {})
        # 按 schema 顺序插入：tblBorders 在 shd/tblLayout/tblCellMar/tblLook 之前
        tblPr.insert_element_before(
            borders, "w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook")
        border(borders, "top", "single", 12)
        border(borders, "bottom", "single", 12)
        for name in ("left", "right", "insideH", "insideV"):
            border(borders, name, "none", 0)
        if has_head:
            for cell in table.rows[0].cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tc_borders = etree.SubElement(tcPr, qn("w:tcBorders"))
                border(tc_borders, "bottom", "single", 6)

    def _fill_table_cell(self, cell, cell_token, bold, align=None):
        p = cell.paragraphs[0]
        align = align or (cell_token.get("attrs") or {}).get("align")
        if align:
            apply_paragraph_format(p, align=align)
        cfg = dict(self.body_cfg)
        if bold:
            cfg["bold"] = True
        self.render_inlines(p, cell_token.get("children", []), cfg)

    def block_block_html(self, doc, token, depth, in_quote):
        pass  # 忽略原始 HTML 块

    def block_block_math(self, doc, token, depth, in_quote):
        # 独立公式 $$...$$：单独成段、居中，用 m:oMathPara 包裹
        from .omml import MathConversionError, latex_to_omml, wrap_omath_para

        p = doc.add_paragraph()
        apply_paragraph_format(p, align="center",
                               space_before=6, space_after=6)
        try:
            omath = latex_to_omml(token.get("raw", ""),
                                  self.body_cfg.get("font_size"))
            p._p.append(wrap_omath_para(omath))
        except MathConversionError:
            self._add_run(p, token.get("raw", ""), self.body_cfg, color="C00000")

    # -------------------------------------------------------------- inlines
    def render_inlines(self, paragraph, children, font_cfg, *,
                       bold=False, italic=False, underline=False, color=None):
        for token in children or []:
            ttype = token.get("type")
            if ttype == "text":
                self._add_run(paragraph, token.get("raw", ""), font_cfg,
                              bold=bold, italic=italic, underline=underline, color=color)
            elif ttype == "strong":
                self.render_inlines(paragraph, token.get("children", []), font_cfg,
                                    bold=True, italic=italic, underline=underline, color=color)
            elif ttype == "emphasis":
                self.render_inlines(paragraph, token.get("children", []), font_cfg,
                                    bold=bold, italic=True, underline=underline, color=color)
            elif ttype == "codespan":
                run = paragraph.add_run(token.get("raw", ""))
                style_run(run, self.code_cfg)
            elif ttype == "link":
                self.render_inlines(paragraph, token.get("children", []), font_cfg,
                                    bold=bold, italic=italic, underline=True,
                                    color="0563C1")
            elif ttype == "image":
                self._add_image(paragraph, token)
            elif ttype in ("softbreak",):
                self._add_run(paragraph, " ", font_cfg, bold=bold, italic=italic,
                              underline=underline, color=color)
            elif ttype == "linebreak":
                paragraph.add_run().add_break()
            elif ttype == "inline_html":
                pass
            elif ttype == "inline_math":
                self._add_math(paragraph, token.get("raw", ""))
            else:
                # 兜底：尝试渲染其子节点或 raw 文本
                if token.get("children"):
                    self.render_inlines(paragraph, token["children"], font_cfg,
                                        bold=bold, italic=italic,
                                        underline=underline, color=color)
                elif token.get("raw"):
                    self._add_run(paragraph, token["raw"], font_cfg, bold=bold,
                                  italic=italic, underline=underline, color=color)

    def _add_run(self, paragraph, text, font_cfg, *, bold=False, italic=False,
                 underline=False, color=None):
        if not text:
            return
        # 参考文献区内的 [N] 拆分为 [ + 数字书签 + ]，便于交叉引用精确指向数字
        if self.in_references:
            m = re.fullmatch(r"\[(\d+)\]", text)
            if m:
                self._add_run(paragraph, "[", font_cfg, bold=bold, italic=italic,
                              underline=underline, color=color)
                run = paragraph.add_run(m.group(1))
                style_run(run, font_cfg,
                          bold=True if bold else font_cfg.get("bold"),
                          italic=True if italic else font_cfg.get("italic"),
                          underline=True if underline else None,
                          color=color or font_cfg.get("color"))
                self._wrap_run_with_bookmark(run, f"ref_{m.group(1)}", self._next_bookmark_id())
                self._add_run(paragraph, "]", font_cfg, bold=bold, italic=italic,
                              underline=underline, color=color)
                return
            run = paragraph.add_run(text)
            style_run(run, font_cfg,
                      bold=True if bold else font_cfg.get("bold"),
                      italic=True if italic else font_cfg.get("italic"),
                      underline=True if underline else None,
                      color=color or font_cfg.get("color"))
            return
        if not self.ref_numbers or not CITATION_RE.search(text):
            run = paragraph.add_run(text)
            style_run(run, font_cfg,
                      bold=True if bold else font_cfg.get("bold"),
                      italic=True if italic else font_cfg.get("italic"),
                      underline=True if underline else None,
                      color=color or font_cfg.get("color"))
            return
        for kind, content in self._split_citation(text):
            if kind == "text":
                run = paragraph.add_run(content)
                style_run(run, font_cfg,
                          bold=True if bold else font_cfg.get("bold"),
                          italic=True if italic else font_cfg.get("italic"),
                          underline=True if underline else None,
                          color=color or font_cfg.get("color"))
            else:
                self._add_citation_links(paragraph, content)

    def _add_math(self, paragraph, latex: str):
        """行内公式 $...$：转为 OMML 挂到当前段落；失败时按原文输出红色文本。"""
        from .omml import MathConversionError, latex_to_omml

        try:
            omath = latex_to_omml(latex, self.body_cfg.get("font_size"))
            paragraph._p.append(omath)
        except MathConversionError:
            self._add_run(paragraph, f"${latex}$", self.body_cfg, color="C00000")

    def _add_image(self, paragraph, token):
        url = (token.get("attrs", {}) or {}).get("url", "")
        if not url or "://" in url:
            alt = "".join(c.get("raw", "") for c in token.get("children", []) or [])
            self._add_run(paragraph, f"[图片: {alt or url}]", self.body_cfg)
            return
        path = url if os.path.isabs(url) else os.path.join(self.base_dir, url)
        if not os.path.exists(path):
            self._add_run(paragraph, f"[图片缺失: {url}]", self.body_cfg, color="C00000")
            return
        run = paragraph.add_run()
        try:
            run.add_picture(path, width=Cm(12))
        except Exception:
            self._add_run(paragraph, f"[图片无法读取: {url}]", self.body_cfg,
                          color="C00000")

    # -------------------------------------------------------------- references / citations

    @staticmethod
    def _inline_text(token) -> str:
        """递归提取 token 中的纯文本。"""
        if not isinstance(token, dict):
            return ""
        if token.get("type") == "text":
            return token.get("raw", "")
        return "".join(Converter._inline_text(child) for child in token.get("children", []) or [])

    def _scan_references(self, tokens):
        """扫描 AST，收集参考文献区中所有引用编号（含 [N] 与 [M-N] 范围）。"""
        in_refs = False
        ref_section_level = 0
        for token in tokens:
            ttype = token.get("type")
            if ttype == "heading":
                text = self._inline_text(token).strip()
                level = int(token.get("attrs", {}).get("level", 1))
                if text in ("参考文献", "References"):
                    in_refs = True
                    ref_section_level = level
                    continue
                if in_refs and level <= ref_section_level:
                    in_refs = False
                    continue
            if not in_refs:
                continue
            if ttype == "list":
                for item in token.get("children", []) or []:
                    self._collect_ref_numbers(item)
            elif ttype in ("paragraph", "block_text"):
                self._collect_ref_numbers(token)

    def _collect_ref_numbers(self, token):
        """从 token 文本中收集 [N] 与 [M-N] 里的所有编号。"""
        text = self._inline_text(token)
        for m in CITATION_RE.finditer(text):
            inner = m.group()[1:-1]
            for seg in inner.split(","):
                seg = seg.strip()
                if "-" in seg:
                    a, b = seg.split("-", 1)
                    try:
                        for num in range(int(a), int(b) + 1):
                            self.ref_numbers.add(num)
                    except ValueError:
                        pass
                else:
                    try:
                        self.ref_numbers.add(int(seg))
                    except ValueError:
                        pass

    def _next_bookmark_id(self) -> int:
        self._bookmark_id += 1
        return self._bookmark_id

    def _add_bookmark(self, paragraph, bookmark_name: str, bookmark_id: int):
        """给段落整体添加书签（用于参考文献条目定位）。"""
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(bookmark_id))
        start.set(qn("w:name"), bookmark_name)
        paragraph._p.insert(0, start)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(bookmark_id))
        paragraph._p.append(end)

    def _wrap_run_with_bookmark(self, run, bookmark_name: str, bookmark_id: int):
        """在段落层级给一个 run 首尾包裹书签（书签标记必须是 w:p 的直接子元素）。"""
        p = run._r.getparent()
        idx = list(p).index(run._r)
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(bookmark_id))
        start.set(qn("w:name"), bookmark_name)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(bookmark_id))
        p.insert(idx, start)
        p.insert(idx + 2, end)

    def _split_citation(self, text: str):
        """把文本拆分为普通文本与引用标注交替的列表。"""
        parts = []
        last = 0
        for m in CITATION_RE.finditer(text):
            if m.start() > last:
                parts.append(("text", text[last:m.start()]))
            parts.append(("cite", m.group()))
            last = m.end()
        if last < len(text):
            parts.append(("text", text[last:]))
        if not parts:
            parts.append(("text", text))
        return parts

    def _add_citation_links(self, paragraph, citation_text: str):
        """把 [1]、[1-2]、[1,3] 等渲染为带中括号的上标，其中数字为交叉引用域。"""
        self._add_sup_run(paragraph, "[")
        inner = citation_text[1:-1]
        segments = [s.strip() for s in inner.split(",")]
        first = True
        for seg in segments:
            if not first:
                self._add_sup_run(paragraph, ",")
            first = False
            if "-" in seg:
                a, b = seg.split("-", 1)
                self._add_cite_field(paragraph, a.strip())
                self._add_sup_run(paragraph, "-")
                self._add_cite_field(paragraph, b.strip())
            else:
                self._add_cite_field(paragraph, seg)
        self._add_sup_run(paragraph, "]")

    def _set_citation_rpr(self, rPr):
        """设置引用标注 run 的公共格式（上标、字体、字号、颜色）。"""
        vert_align = OxmlElement("w:vertAlign")
        vert_align.set(qn("w:val"), "superscript")
        rPr.append(vert_align)

        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), self.body_cfg.get("font_en", ""))
        rfonts.set(qn("w:hAnsi"), self.body_cfg.get("font_en", ""))
        rfonts.set(qn("w:eastAsia"), self.body_cfg.get("font_zh", ""))
        rPr.append(rfonts)

        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(self.body_cfg.get("font_size", 12) * 2)))
        rPr.append(sz)

        color = OxmlElement("w:color")
        color.set(qn("w:val"), "000000")
        rPr.append(color)

    def _add_cite_field(self, paragraph, number_text: str):
        """为单个引用编号创建 Word REF 域交叉引用（带 \\h 开关，可点击跳转）。"""
        try:
            num = int(number_text)
        except ValueError:
            self._add_sup_run(paragraph, number_text)
            return
        if num not in self.ref_numbers:
            self._add_sup_run(paragraph, number_text)
            return

        bookmark_name = f"ref_{num}"

        def _field_run(fld_type=None, text=None, instr=None):
            r = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            self._set_citation_rpr(rPr)
            r.append(rPr)
            if fld_type is not None:
                fld = OxmlElement("w:fldChar")
                fld.set(qn("w:fldCharType"), fld_type)
                r.append(fld)
            if instr is not None:
                el = OxmlElement("w:instrText")
                el.set(qn("xml:space"), "preserve")
                el.text = instr
                r.append(el)
            if text is not None:
                t = OxmlElement("w:t")
                t.text = text
                r.append(t)
            paragraph._p.append(r)

        _field_run(fld_type="begin")
        _field_run(instr=f" REF {bookmark_name} \\h ")
        _field_run(fld_type="separate")
        _field_run(text=number_text)
        _field_run(fld_type="end")

    def _add_sup_run(self, paragraph, text: str):
        """添加一个普通上标 run（用于连接符、逗号或未知编号）。"""
        run = paragraph.add_run(text)
        style_run(run, self.body_cfg)
        run.font.superscript = True

    def _merge_citation_tokens(self, tokens):
        """递归合并被 Mistune 拆开的引用标注文本 token（例如 [ 和 1]）。"""
        for token in tokens:
            if not isinstance(token, dict):
                continue
            children = token.get("children")
            if children:
                token["children"] = self._merge_citation_tokens(children)
                # 仅对 inline 容器执行引用合并
                if token.get("type") in ("text", "strong", "emphasis", "link", "paragraph",
                                          "block_text", "heading", "list_item", "table_cell"):
                    token["children"] = self._merge_inline_citations(token["children"])
        return tokens

    @staticmethod
    def _merge_inline_citations(children):
        """把连续文本 token 中的 [1]、[1-2]、[1,4] 合并成单个文本 token。"""
        if not children:
            return children
        result = []
        i = 0
        n = len(children)
        while i < n:
            child = children[i]
            if child.get("type") == "text" and child.get("raw") == "[":
                j = i + 1
                acc = []
                citation_body = None
                suffix = ""
                while j < n:
                    nxt = children[j]
                    if nxt.get("type") == "text":
                        raw = nxt.get("raw", "")
                        acc.append(raw)
                        combined = "".join(acc)
                        m = re.match(r"^(\d+(?:-\d+)?(?:,\s*\d+(?:-\d+)?)*)\](.*)", combined)
                        if m:
                            citation_body = m.group(1)
                            suffix = m.group(2)
                            break
                        j += 1
                    else:
                        break
                if citation_body is not None:
                    result.append({"type": "text", "raw": f"[{citation_body}]"})
                    if suffix:
                        result.append({"type": "text", "raw": suffix})
                    i = j + 1
                    continue
            result.append(child)
            i += 1
        return result

    @staticmethod
    def _apply_heading_number_sep(token, cfg):
        """把标题开头“数字 + 空格”中的空格替换为配置的间隔符。"""
        sep = cfg.get("number_sep")
        if sep is None or sep == " ":
            return
        children = token.get("children", []) or []
        if not children:
            return
        first = children[0]
        if first.get("type") != "text":
            return
        raw = first.get("raw", "")
        m = re.match(r"^(\d+(?:\.\d+)*)\s+(.*)$", raw)
        if m:
            first["raw"] = m.group(1) + sep + m.group(2)
